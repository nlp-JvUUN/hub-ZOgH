"""
Word 文档（.docx）格式加载器

依赖：python-docx

处理策略：
  - Heading 样式 → title 块
  - 加粗短文本 → title 块（中文文档可能用自定义样式）
  - 原生表格 → markdown table 块
  - 普通段落 → text 块
"""

import re
import logging
from pathlib import Path

from parsed_block_schema import ParsedDocument, ParsedBlock
from .base_loader import table_to_markdown

logger = logging.getLogger(__name__)


class DocxLoader:
    """Word 文档加载器。"""

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in (".docx", ".doc")

    def load(self, file_path: Path, meta: dict) -> ParsedDocument:
        logger.info(f"加载 DOCX: {file_path.name}")

        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "需要安装 python-docx 来加载 .docx 文件:\n"
                "  pip install python-docx"
            )

        doc = DocxDocument(str(file_path))
        blocks = []
        section_stack: list[str] = []

        # python-docx 的 iter_block_items 可以按文档顺序遍历段落和表格
        # 但需要我们自己实现
        from docx.oxml.ns import qn

        body = doc.element.body
        para_idx = 0
        table_idx = 0

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":  # 段落
                para = doc.paragraphs[para_idx]
                para_idx += 1
                text = para.text.strip()
                if not text:
                    continue

                # 判断是否标题
                style_name = (para.style.name if para.style else "").lower()
                is_heading = "heading" in style_name or "标题" in style_name or "title" in style_name

                # 检测加粗
                is_bold = False
                if para.runs:
                    is_bold = any(r.bold for r in para.runs if r.bold)

                if is_heading or (is_bold and len(text) < 80):
                    self._update_section(text, section_stack)
                    blocks.append(ParsedBlock(
                        block_type="title",
                        content=text,
                        page_num=0,
                        section_path=list(section_stack),
                    ))
                else:
                    blocks.append(ParsedBlock(
                        block_type="text",
                        content=text,
                        page_num=0,
                        section_path=list(section_stack),
                    ))

            elif tag == "tbl":  # 表格
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    table_idx += 1

                    rows_data = []
                    for row in table.rows:
                        cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                        rows_data.append(cells)

                    if rows_data:
                        md_table = table_to_markdown(rows_data)
                        blocks.append(ParsedBlock(
                            block_type="table",
                            content=md_table,
                            page_num=0,
                            section_path=list(section_stack),
                        ))

        logger.info(f"  解析完成: {len(blocks)} 个块")
        return ParsedDocument(meta=meta, source=str(file_path), blocks=blocks)

    @staticmethod
    def _update_section(title: str, stack: list[str]):
        """维护章节栈。"""
        if re.match(r"^第[一二三四五六七八九十]+章", title):
            stack.clear()
            stack.append(title)
        elif re.match(r"^第[一二三四五六七八九十]+节", title):
            stack[:] = stack[:1] + [title]
        elif re.match(r"^[一二三四五六七八九十]、", title):
            stack[:] = stack[:2] + [title]
        else:
            stack.append(title)
