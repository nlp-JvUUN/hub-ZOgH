"""
通用文件解析脚本：支持PDF / docx / md / txt 全格式解析
兼容原PDF结构化输出结构，下游分块、向量库、RAG无需修改
PDF能力保留：表格提取、标题层级、页眉噪声过滤、扫描件OCR
Word：段落、标题、表格提取
TXT/MD：纯文本分块，简单标题识别
依赖安装：
pip install pdfplumber pymupdf pytesseract pillow python-docx
# tesseract-ocr 需单独安装配置
"""
import re
import json
import logging
import os
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Optional
import datetime
# PDF解析依赖
import pdfplumber
import fitz
# Word解析依赖
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
# OCR依赖
try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# 目录配置
RAW_DIR    = Path(__file__).parent.parent / "data" / "raw_files"
PARSED_DIR = Path(__file__).parent.parent / "data" / "parsed"
RAW_DIR.mkdir(parents=True, exist_ok=True)
PARSED_DIR.mkdir(parents=True, exist_ok=True)

# OCR路径配置（Windows）
# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# -------------------------- 通用数据结构（和原PDF完全兼容） --------------------------
@dataclass
class ParsedBlock:
    block_type:   str            # "text" | "table" | "title"
    content:      str            # 文本/Markdown表格
    page_num:     int            # PDF页码，非PDF统一填0
    section_path: list[str]      # 章节层级
    is_ocr:       bool = False
    raw_table:    Optional[list] = field(default=None, repr=False)

# -------------------------- 全局工具函数（复用原PDF逻辑） --------------------------
CHAPTER_PATTERNS = [
    re.compile(r"^第[一二三四五六七八九十百]+[章节]"),
    re.compile(r"^[一二三四五六七八九十]、"),
    re.compile(r"^\d+\.\s"),
]
NOISE_PATTERNS = [
    re.compile(r"^.{1,40}年度报告\s*$"),
    re.compile(r"^\d+\s*$"),
    re.compile(r"^—\s*\d+\s*—$"),
]

def is_noise_line(line: str) -> bool:
    line = line.strip()
    if len(line) < 2:
        return True
    return any(p.match(line) for p in NOISE_PATTERNS)

def is_title_line(line: str, fontsize: Optional[float] = None, is_bold: bool = False) -> bool:
    if fontsize and fontsize >= 14:
        return True
    if is_bold and len(line.strip()) < 50:
        return True
    return any(p.match(line.strip()) for p in CHAPTER_PATTERNS)

def table_to_markdown(table: list[list]) -> str:
    if not table:
        return ""
    rows = []
    for row in table:
        cleaned = [str(cell or "").replace("\n", " ").strip() for cell in row]
        rows.append(cleaned)
    if not rows:
        return ""
    header = rows[0]
    lines  = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in rows[1:]:
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[:len(header)]) + " |")
    return "\n".join(lines)

# -------------------------- PDF专用解析器（完全复用原有逻辑） --------------------------
class PdfParser:
    def __init__(self, pdf_path: Path, meta: dict):
        self.pdf_path = pdf_path
        self.meta = meta
        self.blocks: list[ParsedBlock] = []
        self._section_stack: list[str] = []

    def _update_section(self, title: str):
        if re.match(r"^第[一二三四五六七八九十]+章", title):
            self._section_stack = [title]
        elif re.match(r"^第[一二三四五六七八九十]+节", title):
            self._section_stack = self._section_stack[:1] + [title]
        elif re.match(r"^[一二三四五六七八九十]、", title):
            self._section_stack = self._section_stack[:2] + [title]
        else:
            self._section_stack = self._section_stack[:3] + [title]

    def detect_if_scanned(self, page: fitz.Page, text: str) -> bool:
        if len(text.strip()) > 50:
            return False
        image_list = page.get_images(full=True)
        return len(image_list) > 0

    def ocr_page(self, page: fitz.Page, dpi: int = 200) -> str:
        if not OCR_AVAILABLE:
            return "[扫描页，OCR不可用，内容跳过]"
        try:
            mat  = fitz.Matrix(dpi / 72, dpi / 72)
            clip = page.rect
            pix  = page.get_pixmap(matrix=mat, clip=clip)
            img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            text = pytesseract.image_to_string(img, lang="chi_sim+eng")
            return text
        except Exception as e:
            logger.warning(f"OCR失败: {e}")
            return "[扫描页OCR失败]"

    def parse(self) -> list[ParsedBlock]:
        logger.info(f"解析PDF: {self.pdf_path.name}")
        plumber_doc = pdfplumber.open(self.pdf_path)
        fitz_doc    = fitz.open(str(self.pdf_path))
        for page_num in range(len(fitz_doc)):
            fitz_page   = fitz_doc[page_num]
            plumb_page  = plumber_doc.pages[page_num]
            raw_text = fitz_page.get_text("text")
            is_scanned = self.detect_if_scanned(fitz_page, raw_text)
            if is_scanned:
                ocr_text = self.ocr_page(fitz_page)
                self.blocks.append(ParsedBlock(
                    block_type="text",
                    content=ocr_text,
                    page_num=page_num + 1,
                    section_path=list(self._section_stack),
                    is_ocr=True,
                ))
                continue
            # 提取表格
            for table in plumb_page.extract_tables():
                if table:
                    md = table_to_markdown(table)
                    self.blocks.append(ParsedBlock(
                        block_type="table",
                        content=md,
                        page_num=page_num + 1,
                        section_path=list(self._section_stack),
                        raw_table=table,
                    ))
            # 文字+标题解析
            page_dict = fitz_page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            current_para_lines = []
            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    line_text = "".join(span["text"] for span in line.get("spans", [])).strip()
                    if not line_text or is_noise_line(line_text):
                        continue
                    spans    = line.get("spans", [])
                    fontsize = spans[0].get("size", 0) if spans else 0
                    is_bold  = any("Bold" in span.get("font", "") for span in spans)
                    if is_title_line(line_text, fontsize, is_bold):
                        if current_para_lines:
                            self.blocks.append(ParsedBlock(
                                block_type="text",
                                content="\n".join(current_para_lines),
                                page_num=page_num + 1,
                                section_path=list(self._section_stack),
                            ))
                            current_para_lines = []
                        self._update_section(line_text)
                        self.blocks.append(ParsedBlock(
                            block_type="title",
                            content=line_text,
                            page_num=page_num + 1,
                            section_path=list(self._section_stack),
                        ))
                    else:
                        current_para_lines.append(line_text)
            if current_para_lines:
                self.blocks.append(ParsedBlock(
                    block_type="text",
                    content="\n".join(current_para_lines),
                    page_num=page_num + 1,
                    section_path=list(self._section_stack),
                ))
        plumber_doc.close()
        fitz_doc.close()
        logger.info(f"PDF解析完成，共{len(self.blocks)}个块")
        return self.blocks

# -------------------------- Word(docx)解析器 --------------------------
class DocxParser:
    def __init__(self, docx_path: Path, meta: dict):
        self.docx_path = docx_path
        self.meta = meta
        self.blocks: list[ParsedBlock] = []
        self._section_stack = []

    def parse(self) -> list[ParsedBlock]:
        if not DOCX_AVAILABLE:
            logger.error("未安装python-docx，跳过解析")
            return []
        logger.info(f"解析Word: {self.docx_path.name}")
        doc = Document(str(self.docx_path))
        # 段落与标题
        current_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 判断标题
            if para.style.name.startswith("Heading") or is_title_line(text, is_bold=True):
                if current_text:
                    self.blocks.append(ParsedBlock(
                        block_type="text",
                        content="\n".join(current_text),
                        page_num=0,
                        section_path=list(self._section_stack),
                    ))
                    current_text = []
                self._section_stack.append(text)
                self.blocks.append(ParsedBlock(
                    block_type="title",
                    content=text,
                    page_num=0,
                    section_path=list(self._section_stack),
                ))
            else:
                current_text.append(text)
        if current_text:
            self.blocks.append(ParsedBlock(
                block_type="text",
                content="\n".join(current_text),
                page_num=0,
                section_path=list(self._section_stack),
            ))
        # 表格
        for table in doc.tables:
            raw_table = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                raw_table.append(row_data)
            md_table = table_to_markdown(raw_table)
            self.blocks.append(ParsedBlock(
                block_type="table",
                content=md_table,
                page_num=0,
                section_path=list(self._section_stack),
                raw_table=raw_table
            ))
        logger.info(f"Word解析完成，共{len(self.blocks)}个块")
        return self.blocks

# -------------------------- TXT / Markdown 纯文本解析器 --------------------------
class TextParser:
    def __init__(self, file_path: Path, meta: dict):
        self.file_path = file_path
        self.meta = meta
        self.blocks: list[ParsedBlock] = []
        self._section_stack = []

    def parse(self) -> list[ParsedBlock]:
        logger.info(f"解析文本文件: {self.file_path.name}")
        with open(self.file_path, "r", encoding="utf-8", errors="ignore") as f:
            lines = [line.strip() for line in f.readlines() if line.strip()]
        buffer = []
        for line in lines:
            # Markdown标题 # / ##
            if line.startswith("#"):
                if buffer:
                    self.blocks.append(ParsedBlock(
                        block_type="text",
                        content="\n".join(buffer),
                        page_num=0,
                        section_path=list(self._section_stack),
                    ))
                    buffer = []
                title_text = line.lstrip("# ").strip()
                self._section_stack.append(title_text)
                self.blocks.append(ParsedBlock(
                    block_type="title",
                    content=title_text,
                    page_num=0,
                    section_path=list(self._section_stack),
                ))
            elif is_title_line(line):
                if buffer:
                    self.blocks.append(ParsedBlock(
                        block_type="text",
                        content="\n".join(buffer),
                        page_num=0,
                        section_path=list(self._section_stack),
                    ))
                    buffer = []
                self._section_stack.append(line)
                self.blocks.append(ParsedBlock(
                    block_type="title",
                    content=line,
                    page_num=0,
                    section_path=list(self._section_stack),
                ))
            else:
                buffer.append(line)
        if buffer:
            self.blocks.append(ParsedBlock(
                block_type="text",
                content="\n".join(buffer),
                page_num=0,
                section_path=list(self._section_stack),
            ))
        logger.info(f"文本文件解析完成，共{len(self.blocks)}个块")
        return self.blocks

# -------------------------- 统一分发入口 --------------------------
def parse_single_file(file_path: Path) -> dict:
    """自动识别文件类型，分发对应解析器，输出统一结构化json数据"""
    suffix = file_path.suffix.lower()
    # 通用基础元信息
    stat = file_path.stat()
    base_meta = {
        "filename": file_path.name,
        "file_suffix": suffix,
        "source_path": str(file_path.resolve()),
        "file_size_kb": round(stat.st_size / 1024, 2),
        "create_time": datetime.datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
        "stock_code": "unknown",
        "year": "unknown",
    }
    blocks = []
    if suffix == ".pdf":
        parser = PdfParser(file_path, base_meta)
        blocks = parser.parse()
    elif suffix == ".docx":
        parser = DocxParser(file_path, base_meta)
        blocks = parser.parse()
    elif suffix in (".txt", ".md"):
        parser = TextParser(file_path, base_meta)
        blocks = parser.parse()
    else:
        logger.warning(f"不支持的文件格式 {suffix}，跳过: {file_path.name}")
        return None
    output_data = {
        "meta": base_meta,
        "source": str(file_path),
        "blocks": [asdict(b) for b in blocks],
    }
    # 保存解析结果
    out_path = PARSED_DIR / f"{file_path.stem}_{suffix.strip('.')}.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output_data, f, ensure_ascii=False, indent=2)
    logger.info(f"解析结果保存至: {out_path}")
    return output_data

# -------------------------- 主流程 --------------------------
def main():
    # 遍历raw_files下所有支持文件
    support_suffix = {".pdf", ".docx", ".txt", ".md"}
    all_files = []
    for suf in support_suffix:
        all_files.extend(RAW_DIR.glob(f"*{suf}"))
    if not all_files:
        logger.error(f"未找到可解析文件，请将文件放入 {RAW_DIR}")
        return
    logger.info(f"共检测到 {len(all_files)} 个待解析文件")
    for file in all_files:
        parse_single_file(file)
    logger.info(f"\n全部文件解析完成，解析结果存放于 {PARSED_DIR}")

if __name__ == "__main__":
    main()