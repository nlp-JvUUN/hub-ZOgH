# -*- coding: utf-8 -*-
"""Convert a template-style resume Markdown into PDF / HTML / Word / Markdown.

Portable: no hardcoded machine paths. A CJK font is auto-detected from common
locations on Windows / macOS / Linux, or supplied via --font / --font-dir.
Layout is driven by styles.json (presets: classic / modern), overridable with
--style or a custom styles file.

Usage:
    python md_to_pdf.py <resume.md> [output] [--format pdf|html|docx|md]
                       [--style classic|modern] [--styles <styles.json>]
                       [--font <path> | --font-dir <dir>]

Input is the Markdown produced by the resume-from-template skill, i.e. the
structure of the bundled template (templates/resume-template.md):
    # 姓名                  -> H1
    **性别 | 年龄：X岁 | ...**  -> contact line
    **X年工作经验 | 求职意向 ...** -> tagline line
    ## 章节                 -> H2
    ### 公司　职位　起止时间   -> H3
    - 条目                   -> bullet (grouped into items)
    内容：/ 业绩：            -> plain label line
    <!-- 待确认项 ... -->    -> skipped
"""
import argparse
import io
import json
import re
import sys
from pathlib import Path

import fitz

BULLET = "●"  # ● — SimHei has this glyph; U+2022 '•' does not.
MM_PER_PT = 25.4 / 72

SCRIPT_DIR = Path(__file__).resolve().parent
STYLES_FILE = SCRIPT_DIR / "styles.json"

# Fonts are searched in this order; the first existing one wins.
CJK_FONT_CANDIDATES = [
    # Windows
    r"C:/Windows/Fonts/simhei.ttf",
    r"C:/Windows/Fonts/msyh.ttc",
    r"C:/Windows/Fonts/msyh.ttf",
    r"C:/Windows/Fonts/simsun.ttc",
    r"C:/Windows/Fonts/Deng.ttf",
    # macOS
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
    # Linux
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
]


def load_styles(style_name=None, styles_path=None):
    """Load a style dict. --styles <file> takes precedence over bundled styles.json."""
    path = Path(styles_path) if styles_path else STYLES_FILE
    if not path.is_file():
        raise SystemExit(f"Styles file not found: {path}")
    data = json.loads(path.read_text(encoding="utf-8"))
    styles = data.get("styles", {})
    name = style_name or data.get("default_style", "classic")
    if name not in styles:
        raise SystemExit(
            f"Unknown style: {name}. Available: {', '.join(styles)}")
    cfg = styles[name]
    cfg.setdefault("spacing", {})
    return cfg


def build_css(cfg, font_basename):
    mm = lambda v: f"{v}mm"
    s = cfg["spacing"]
    css = f"""
@font-face {{ font-family: "zh"; src: url("{font_basename}"); }}
* {{ box-sizing: border-box; }}
body {{
  font-family: "zh", sans-serif;
  font-size: {cfg['font_size_pt']}pt;
  color: {cfg['color_body']};
  line-height: {cfg['line_height']};
  margin: {mm(cfg['margin_top_mm'])} {mm(cfg['margin_right_mm'])} {mm(cfg['margin_bottom_mm'])} {mm(cfg['margin_left_mm'])};
}}
h1 {{ font-size: {cfg['h1_size_pt']}pt; margin: 0 0 {s.get('h1_mb_pt',3)}pt 0; color: {cfg['color_name']}; }}
.contact {{ font-size: {cfg['contact_size_pt']}pt; color: {cfg['color_contact']}; margin-bottom: {s.get('contact_mb_pt',3)}pt; }}
.tagline {{ font-size: {cfg['tagline_size_pt']}pt; color: {cfg['color_contact']}; margin-bottom: {s.get('tagline_mb_pt',8)}pt; }}
h2 {{
  font-size: {cfg['h2_size_pt']}pt;
  color: {cfg['color_h2']};
  border-bottom: {cfg['h2_border_bottom_pt']}pt solid {cfg['color_h2']};
  padding-bottom: 2pt;
  margin: {s.get('h2_mt_pt',9)}pt 0 {s.get('h2_mb_pt',4)}pt 0;
  page-break-after: avoid;
}}
h3 {{ font-size: {cfg['h3_size_pt']}pt; color: {cfg['color_body']}; margin: {s.get('h3_mt_pt',7)}pt 0 {s.get('h3_mb_pt',2)}pt 0; page-break-after: avoid; }}
.items {{ margin: {s.get('items_mt_pt',1)}pt 0 {s.get('items_mb_pt',3)}pt 0; }}
.item {{ margin: 0 0 {s.get('item_mb_pt',2.4)}pt 0; padding-left: {s.get('item_indent_pt',14)}pt; text-indent: -{s.get('item_indent_pt',14)}pt; }}
.bullet {{ font-size: {cfg['bullet_em']}em; vertical-align: middle; }}
.sub {{ color: {cfg['color_sub']}; font-size: {cfg['sub_size_pt']}pt; margin-bottom: {s.get('sub_mb_pt',3)}pt; }}
.edu {{ margin-top: {s.get('edu_mt_pt',3)}pt; }}
.plain {{ margin-bottom: {s.get('plain_mb_pt',3)}pt; }}
"""
    return css


def find_cjk_font(font=None, font_dir=None):
    """Resolve a CJK font file path. Priority: --font > --font-dir > auto-detect."""
    if font:
        path = Path(font)
        if not path.is_file():
            raise SystemExit(f"Font not found: {font}")
        return path

    if font_dir:
        d = Path(font_dir)
        if not d.is_dir():
            raise SystemExit(f"Font dir not found: {font_dir}")
        for name in ("simhei.ttf", "msyh.ttc", "msyh.ttf", "simsun.ttc",
                     "Deng.ttf", "NotoSansCJK-Regular.ttc", "wqy-zenhei.ttc",
                     "wqy-microhei.ttc", "PingFang.ttc"):
            p = d / name
            if p.is_file():
                return p
        raise SystemExit(
            f"No CJK font found in {font_dir}. Pass --font <path> explicitly.")

    for candidate in CJK_FONT_CANDIDATES:
        if Path(candidate).is_file():
            return Path(candidate)
    raise SystemExit(
        "No CJK font auto-detected. Pass --font <path-to-a-CJK-font-file>.")


def escape(text):
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def inline(text):
    """Escape then convert **bold** markers to HTML <b> tags."""
    text = escape(text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    return text


def md_to_html(md, css):
    body = []
    in_list = False
    in_comment = False

    def close_list():
        nonlocal in_list
        if in_list:
            body.append("</div>")
            in_list = False

    for raw in md.splitlines():
        line = raw.strip()

        if in_comment:
            if "-->" in line:
                in_comment = False
            continue
        if line.startswith("<!--"):
            if "-->" not in line:
                in_comment = True
            continue

        if not line:
            continue
        if line == "---":
            close_list()
            continue

        m = re.match(r"^# (.*)$", line)
        if m:
            close_list()
            body.append(f"<h1>{inline(m.group(1))}</h1>")
            continue

        m = re.match(r"^## (.*)$", line)
        if m:
            close_list()
            body.append(f"<h2>{inline(m.group(1))}</h2>")
            continue

        m = re.match(r"^### (.*)$", line)
        if m:
            close_list()
            body.append(f"<h3>{inline(m.group(1))}</h3>")
            continue

        m = re.match(r"^\s*[-*] (.*)$", line)
        if m:
            if not in_list:
                body.append('<div class="items">')
                in_list = True
            body.append(f'<div class="item"><span class="bullet">{BULLET}</span> {inline(m.group(1))}</div>')
            continue

        # Bold-only lines: contact (first) vs tagline (contains 工作经验/求职意向)
        if line.startswith("**") and line.endswith("**"):
            cls = "tagline" if ("工作经验" in line or "求职意向" in line) else "contact"
            close_list()
            body.append(f'<div class="{cls}">{inline(line)}</div>')
            continue

        close_list()
        body.append(f'<div class="plain">{inline(line)}</div>')

    close_list()
    return f"<html><head><meta charset='utf-8'><style>{css}</style></head><body>{''.join(body)}</body></html>"


def find_leftover_placeholders(md):
    """Return all unresolved {{...}} / }} tokens found in the source markdown."""
    found = sorted(set(re.findall(r"\{\{[^{}]*\}\}|\{\{|\}\}", md)))
    return found


# ---------------------------------------------------------------- exporters
def export_pdf(html, cfg, font_path, out_path):
    buf = io.BytesIO()
    writer = fitz.DocumentWriter(buf)
    archive = fitz.Archive(str(font_path.parent))
    story = fitz.Story(html=html, archive=archive)
    mediabox = fitz.paper_rect(cfg.get("paper", "a4"))
    pages = 0
    while pages < 10:
        dev = writer.begin_page(mediabox)
        filled, _ = story.place(mediabox)
        story.draw(dev)
        writer.end_page()
        pages += 1
        if filled != 1:
            break
    writer.close()

    pdf = fitz.open(stream=buf.getvalue(), filetype="pdf")
    pdf.subset_fonts(fallback=True)

    if pdf.page_count > 1:
        n = pdf.page_count
        gray = (0.35, 0.35, 0.35)
        for i, page in enumerate(pdf):
            w = page.rect.width
            label = f"{i + 1} / {n}"
            tw = fitz.get_text_length(label, fontsize=8, fontname="helv")
            page.insert_text(
                ((w - tw) / 2, page.rect.height - 18),
                label, fontsize=8, fontname="helv", color=gray)

    pdf.save(str(out_path), garbage=4, deflate=True)
    pdf.close()
    return pages


def export_html(html, out_path):
    out_path.write_text(html, encoding="utf-8")


def export_docx(md, cfg, font_path, out_path):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    font_name = font_path.stem if font_path.suffix.lower() == ".ttf" else "SimHei"

    def set_font(run, size_pt, bold=False, color=None, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font_name)
        rfonts.set(qn("w:ascii"), font_name)
        rfonts.set(qn("w:hAnsi"), font_name)

    def add_bottom_border(p, color_hex, sz="8"):
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), sz)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex.lstrip("#"))
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_runs(p, text, size_pt, bold=False, color=None):
        for seg in re.split(r"(\*\*.+?\*\*)", text):
            if not seg:
                continue
            if seg.startswith("**") and seg.endswith("**"):
                r = p.add_run(seg[2:-2])
                set_font(r, size_pt, bold=True, color=color)
            else:
                r = p.add_run(seg)
                set_font(r, size_pt, bold=bold, color=color)

    s = cfg["spacing"]
    doc = Document()
    # tighten default margins to match the theme
    for section in doc.sections:
        section.top_margin = Pt(cfg["margin_top_mm"] / MM_PER_PT)
        section.bottom_margin = Pt(cfg["margin_bottom_mm"] / MM_PER_PT)
        section.left_margin = Pt(cfg["margin_left_mm"] / MM_PER_PT)
        section.right_margin = Pt(cfg["margin_right_mm"] / MM_PER_PT)

    in_list = False
    in_comment = False
    for raw in md.splitlines():
        line = raw.strip()
        if in_comment:
            if "-->" in line:
                in_comment = False
            continue
        if line.startswith("<!--"):
            if "-->" not in line:
                in_comment = True
            continue
        if not line:
            continue
        if line == "---":
            in_list = False
            continue

        m = re.match(r"^# (.*)$", line)
        if m:
            in_list = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(s.get("h1_mb_pt", 3))
            add_runs(p, m.group(1), cfg["h1_size_pt"], bold=True, color=cfg["color_name"])
            continue

        m = re.match(r"^## (.*)$", line)
        if m:
            in_list = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(s.get("h2_mt_pt", 9))
            p.paragraph_format.space_after = Pt(s.get("h2_mb_pt", 4))
            add_runs(p, m.group(1), cfg["h2_size_pt"], bold=True, color=cfg["color_h2"])
            add_bottom_border(p, cfg["color_h2"], sz=str(int(cfg["h2_border_bottom_pt"] * 8)))
            continue

        m = re.match(r"^### (.*)$", line)
        if m:
            in_list = False
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(s.get("h3_mt_pt", 7))
            p.paragraph_format.space_after = Pt(s.get("h3_mb_pt", 2))
            add_runs(p, m.group(1), cfg["h3_size_pt"], bold=True, color=cfg["color_body"])
            continue

        m = re.match(r"^\s*[-*] (.*)$", line)
        if m:
            if not in_list:
                in_list = True
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Pt(s.get("item_indent_pt", 14))
            pf.first_line_indent = Pt(-s.get("item_indent_pt", 14))
            pf.space_after = Pt(s.get("item_mb_pt", 2.4))
            r = p.add_run(f"{BULLET} ")
            set_font(r, cfg["font_size_pt"], color=cfg["color_h2"])
            add_runs(p, m.group(1), cfg["font_size_pt"])
            continue

        if line.startswith("**") and line.endswith("**"):
            in_list = False
            is_tagline = "工作经验" in line or "求职意向" in line
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(
                s.get("tagline_mb_pt", 8) if is_tagline else s.get("contact_mb_pt", 3))
            add_runs(p, line, cfg["tagline_size_pt"] if is_tagline else cfg["contact_size_pt"],
                     color=cfg["color_contact"])
            continue

        in_list = False
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(s.get("plain_mb_pt", 3))
        add_runs(p, line, cfg["font_size_pt"])

    doc.save(str(out_path))


def export_md(md, out_path):
    out_path.write_text(md, encoding="utf-8")


# ---------------------------------------------------------------- main
def main():
    ap = argparse.ArgumentParser(
        description="Convert resume Markdown to PDF / HTML / Word / Markdown.")
    ap.add_argument("input", help="Path to the resume Markdown file")
    ap.add_argument("output", nargs="?", help="Output path (default: <input>.<format>)")
    ap.add_argument("--format", choices=["pdf", "html", "docx", "md"],
                    help="Output format (default: inferred from output extension, else pdf)")
    ap.add_argument("--style", help="Style preset name from styles.json")
    ap.add_argument("--styles", help="Path to a custom styles.json")
    ap.add_argument("--font", help="Explicit CJK font file path (e.g. simhei.ttf)")
    ap.add_argument("--font-dir", help="Directory to search for a CJK font")
    ap.add_argument("--check-only", action="store_true",
                    help="Validate the markdown (placeholders) without generating output")
    args = ap.parse_args()

    md_path = Path(args.input)
    md = md_path.read_text(encoding="utf-8")

    # Self-validation: unresolved placeholders are the #1 mistake in filled resumes.
    leftover = find_leftover_placeholders(md)
    if leftover:
        print(f"ERROR: unresolved placeholder(s) found: {leftover}", file=sys.stderr)
        print("Replace every {{...}} with real content before export.", file=sys.stderr)
        sys.exit(2)
    if args.check_only:
        print(f"OK: {md_path} has no unresolved placeholders")
        return

    cfg = load_styles(args.style, args.styles)

    # Resolve output path + format
    out = Path(args.output) if args.output else md_path.with_suffix("")
    fmt = args.format
    if fmt is None:
        if out.suffix.lower() in (".pdf", ".html", ".docx", ".md"):
            fmt = out.suffix.lower().lstrip(".")
        else:
            fmt = "pdf"
    if out.suffix.lower() not in (".pdf", ".html", ".docx", ".md"):
        out = out.with_suffix("." + fmt)

    if fmt == "md":
        export_md(md, out)
        print(f"Markdown saved: {out}")
        return

    font_path = find_cjk_font(args.font, args.font_dir)
    css = build_css(cfg, font_path.name)
    html = md_to_html(md, css)

    if fmt == "html":
        export_html(html, out)
        print(f"HTML saved: {out}")
    elif fmt == "docx":
        export_docx(md, cfg, font_path, out)
        print(f"Word saved: {out}")
    else:
        pages = export_pdf(html, cfg, font_path, out)
        print(f"PDF saved: {out} ({pages} pages)")


if __name__ == "__main__":
    main()
