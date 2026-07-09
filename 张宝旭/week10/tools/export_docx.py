# -*- coding: utf-8 -*-
"""
将 kb/entries.json 导出成 docx。

- 文档顶部：Title
- 每条 entry：Heading 2 标题 + 正文（可选：标签）
- 正文支持：HTML（富文本）或 blocks（纯文本）
"""
from __future__ import annotations

import io
import json
import re
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches

ROOT = Path(__file__).resolve().parent.parent
KB_DIR = ROOT / "kb"


def kb_paths(kb_id: str | None = None):
    """返回 (entries_json_path, images_dir, kb_name)。"""
    # 优先读 kb/index.json 拿 default 库 + 显示名
    idx_path = KB_DIR / "index.json"
    name = None
    target = kb_id
    if idx_path.exists():
        try:
            idx = json.loads(idx_path.read_text(encoding="utf-8"))
            if not target:
                target = idx.get("default")
            for k in idx.get("kbs", []):
                if k["id"] == target:
                    name = k.get("name") or k["id"]
                    break
        except Exception:
            pass
    # 默认兜底：默认目录是 kb/
    if not target:
        # 旧布局：kb/entries.json
        if (KB_DIR / "entries.json").exists():
            return KB_DIR / "entries.json", KB_DIR / "images", "知识库"
        target = "default"
    if not name:
        name = target
    return KB_DIR / target / "entries.json", KB_DIR / target / "images", name


def add_meta_line(doc: Document, entry: dict):
    """只在有标签时写一行小字。状态/来源都不再导出。"""
    if not entry.get("tags"):
        return
    p = doc.add_paragraph()
    run = p.add_run("标签：" + " / ".join(entry["tags"]))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# --------- HTML -> docx 简易转换 ---------
class HtmlToDocx(HTMLParser):
    """把 Quill 输出的 HTML 写进 python-docx。
    支持：p / h1~h4 / ul/ol/li / strong-b / em-i / u / br / img / table。
    其他标签按"忽略标签、保留文字"处理。
    """

    def __init__(self, doc: Document, image_root: Path):
        super().__init__(convert_charrefs=True)
        self.doc = doc
        self.image_root = image_root
        self.stack: list[str] = []
        self.bold = 0
        self.italic = 0
        self.underline = 0
        self.list_kind: list[str] = []
        self.cur_para = None  # docx Paragraph
        self.in_table = False
        self.table = None
        self.table_rows = []
        self.cur_row = None
        self.cur_cell_text = ""

    # ---- 工具 ----
    def _ensure_para(self):
        if self.cur_para is None:
            self.cur_para = self.doc.add_paragraph()
        return self.cur_para

    def _new_para(self, style=None):
        self.cur_para = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()
        return self.cur_para

    def _add_run(self, text: str):
        if not text:
            return
        if self.in_table:
            self.cur_cell_text += text
            return
        p = self._ensure_para()
        run = p.add_run(text)
        if self.bold: run.bold = True
        if self.italic: run.italic = True
        if self.underline: run.underline = True

    # ---- HTMLParser hooks ----
    def handle_starttag(self, tag, attrs):
        attrs = dict(attrs)
        self.stack.append(tag)
        if tag in ("p", "div"):
            self._new_para()
        elif tag in ("h1", "h2", "h3", "h4"):
            level = int(tag[1])
            # entry 自身的标题已经写过了，这里 H1/H2 都降级成 Heading 4，避免破坏整体层级
            style = "Heading 4" if level <= 2 else f"Heading {min(level, 4)}"
            self._new_para(style=style)
        elif tag in ("strong", "b"):
            self.bold += 1
        elif tag in ("em", "i"):
            self.italic += 1
        elif tag == "u":
            self.underline += 1
        elif tag == "br":
            self._add_run("\n")
        elif tag in ("ul", "ol"):
            self.list_kind.append(tag)
        elif tag == "li":
            style = "List Bullet" if (self.list_kind and self.list_kind[-1] == "ul") else "List Number"
            self._new_para(style=style)
        elif tag == "img":
            src = attrs.get("src", "")
            self._add_image(src)
        elif tag == "table":
            self.in_table = True
            self.table_rows = []
            self.cur_row = None
            self.cur_cell_text = ""
        elif tag == "tr":
            self.cur_row = []
        elif tag in ("td", "th"):
            self.cur_cell_text = ""

    def handle_endtag(self, tag):
        if self.stack and self.stack[-1] == tag:
            self.stack.pop()
        if tag in ("strong", "b"):
            self.bold = max(0, self.bold - 1)
        elif tag in ("em", "i"):
            self.italic = max(0, self.italic - 1)
        elif tag == "u":
            self.underline = max(0, self.underline - 1)
        elif tag in ("ul", "ol"):
            if self.list_kind and self.list_kind[-1] == tag:
                self.list_kind.pop()
        elif tag in ("td", "th"):
            if self.cur_row is not None:
                self.cur_row.append(self.cur_cell_text)
            self.cur_cell_text = ""
        elif tag == "tr":
            if self.cur_row is not None:
                self.table_rows.append(self.cur_row)
            self.cur_row = None
        elif tag == "table":
            self._flush_table()
            self.in_table = False
            self.table_rows = []

    def handle_data(self, data: str):
        if not data:
            return
        # 在 li/p/h 内部时直接 add；否则没段落就开一个
        self._add_run(data)

    def _add_image(self, src: str):
        # 仅支持本地 /kb/images/xxx 或 images/xxx
        if not src:
            return
        m = re.match(r"^(?:/kb/)?images/(.+)$", src)
        if not m:
            return
        path = self.image_root / m.group(1)
        if not path.exists():
            return
        if self.in_table:
            self.cur_cell_text += "[图]"
            return
        p = self._ensure_para()
        try:
            run = p.add_run()
            run.add_picture(str(path), width=Inches(5.0))
        except Exception:
            pass
        # 图片单独一段，下个文字再开新段
        self.cur_para = None

    def _flush_table(self):
        rows = [r for r in self.table_rows if r]
        if not rows:
            return
        cols = max(len(r) for r in rows)
        rows = [r + [""] * (cols - len(r)) for r in rows]
        tbl = self.doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Light Grid"
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                tbl.cell(i, j).text = cell_text


def render_blocks(doc: Document, blocks: list, image_root: Path):
    for b in blocks or []:
        t = b.get("type")
        if t == "h4":
            doc.add_paragraph(b.get("text", ""), style="Heading 4")
        elif t == "p":
            doc.add_paragraph(b.get("text", ""))
        elif t == "img":
            src = b.get("src", "")
            m = re.match(r"^(?:/kb/)?images/(.+)$", src)
            if not m:
                continue
            path = image_root / m.group(1)
            if path.exists():
                try:
                    doc.add_picture(str(path), width=Inches(5.0))
                except Exception:
                    pass
        elif t == "table":
            rows = b.get("rows") or []
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            rows = [r + [""] * (cols - len(r)) for r in rows]
            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = "Light Grid"
            for i, row in enumerate(rows):
                for j, c in enumerate(row):
                    tbl.cell(i, j).text = c


def export(out_path: Path | None = None, kb_id: str | None = None) -> Path:
    entries_json, img_dir, kb_name = kb_paths(kb_id)
    if not entries_json.exists():
        raise SystemExit(f"找不到 {entries_json}")
    entries = json.loads(entries_json.read_text(encoding="utf-8"))

    doc = Document()
    title = doc.add_heading(kb_name, level=0)
    note = doc.add_paragraph()
    note_run = note.add_run(f"导出自知识助手 · 共 {len(entries)} 条")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for e in entries:
        doc.add_heading(e.get("title", "(无标题)"), level=2)
        add_meta_line(doc, e)
        if e.get("html"):
            parser = HtmlToDocx(doc, img_dir)
            parser.feed(e["html"])
            parser.cur_para = None
        else:
            render_blocks(doc, e.get("blocks"), img_dir)
        doc.add_paragraph("")

    if out_path is None:
        out_path = ROOT / f"{kb_name}-export.docx"
    doc.save(str(out_path))
    return out_path


def export_to_bytes(kb_id: str | None = None) -> bytes:
    entries_json, img_dir, kb_name = kb_paths(kb_id)
    if not entries_json.exists():
        raise SystemExit(f"找不到 {entries_json}")
    entries = json.loads(entries_json.read_text(encoding="utf-8"))
    doc = Document()
    doc.add_heading(kb_name, level=0)
    note = doc.add_paragraph()
    note_run = note.add_run(f"导出自知识助手 · 共 {len(entries)} 条")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    for e in entries:
        doc.add_heading(e.get("title", "(无标题)"), level=2)
        add_meta_line(doc, e)
        if e.get("html"):
            parser = HtmlToDocx(doc, img_dir)
            parser.feed(e["html"])
        else:
            render_blocks(doc, e.get("blocks"), img_dir)
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    import sys
    kb = sys.argv[1] if len(sys.argv) > 1 else None
    out = export(kb_id=kb)
    print(f"导出完成：{out}")
